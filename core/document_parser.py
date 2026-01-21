import re
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass

@dataclass
class ParsedDocument:
    content: str
    filename: str
    extension: str
    size: int
    raw_text: str

class DocumentParser:
    """Parse multiple document formats."""

    SUPPORTED = {'.txt', '.md', '.afl', '.csv', '.json', '.pdf', '.docx', '.html', '.htm'}

    @classmethod
    def parse(cls, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in cls.SUPPORTED:
            raise ValueError(f"Unsupported: {ext}. Supported: {cls.SUPPORTED}")

        content = cls._extract(path, ext)
        content = cls._clean(content)

        return ParsedDocument(
            content=content,
            filename=path.name,
            extension=ext,
            size=path.stat().st_size,
            raw_text=content
        )

    @classmethod
    def _extract(cls, path: Path, ext: str) -> str:
        if ext in {'.txt', '.md', '.afl', '.csv', '.html', '.htm'}:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if ext in {'.html', '.htm'}:
                text = re.sub(r'<[^>]+>', ' ', text)
            return text

        elif ext == '.json':
            import json
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return cls._flatten_json(data)

        elif ext == '.pdf':
            return cls._extract_pdf(path)

        elif ext == '.docx':
            return cls._extract_docx(path)

        return ""

    @classmethod
    def _extract_pdf(cls, path: Path) -> str:
        try:
            import fitz
            doc = fitz.open(path)
            text = "".join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                return "[PDF parsing unavailable]"

    @classmethod
    def _extract_docx(cls, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return "[DOCX parsing unavailable]"

    @classmethod
    def _flatten_json(cls, data, prefix="") -> str:
        lines = []
        if isinstance(data, dict):
            for k, v in data.items():
                lines.append(cls._flatten_json(v, f"{prefix}{k}: "))
        elif isinstance(data, list):
            for i, item in enumerate(data):
                lines.append(cls._flatten_json(item, f"{prefix}[{i}] "))
        else:
            lines.append(f"{prefix}{data}")
        return "\n".join(lines)

    @classmethod
    def _clean(cls, content: str) -> str:
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = re.sub(r' {2,}', ' ', content)
        content = re.sub(r'Page \d+ of \d+', '', content)
        return content.strip()

    @classmethod
    def parse_batch(cls, paths: List[str]) -> List[ParsedDocument]:
        results = []
        for p in paths:
            try:
                results.append(cls.parse(p))
            except Exception as e:
                print(f"Parse error {p}: {e}")
        return results
